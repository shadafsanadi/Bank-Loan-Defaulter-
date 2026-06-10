"""
Build the Home Loan Default — Interview Prep Study Guide (Word .docx).

This script is the single source of truth for the study guide. Each day of the
30-day plan, content is appended here and the script is re-run to regenerate
study/Interview_Prep_Study_Guide.docx from scratch (idempotent).

Run:  C:\\Users\\Admin\\anaconda3\\python.exe study\\build_study_guide.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ─────────────────────────────────────────────────────────────────────────────
# Styling constants
# ─────────────────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)

OUT_PATH = Path(__file__).parent / "Interview_Prep_Study_Guide.docx"


# ─────────────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────────────
def setup_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Home Loan Default Prediction")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Interview Preparation — Study Guide")
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A 30-Day End-to-End Deep Dive")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Dataset: Home Credit Default Risk  •  Best model: XGBoost (Test ROC-AUC 0.7864)\n"
        "Built with: Python, scikit-learn, XGBoost, FastAPI, SHAP, Docker"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY
    # bottom border
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    # Support simple **bold** segments
    _add_runs_with_bold(p, text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    _add_runs_with_bold(p, text)
    return p


def _add_runs_with_bold(p, text):
    """Render a string supporting **bold** markers."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:
            run.font.bold = True


def mono(doc, text):
    """Monospace block for diagrams / code (preserves spacing)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = NAVY
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# CONTENT — DAY 1
# ─────────────────────────────────────────────────────────────────────────────
def day1_business(doc):
    h1(doc, "Day 1 — Session 1: The Business Problem")
    para(doc, "Taught from the perspective of a Senior Banking Risk Analyst. "
              "Goal: understand WHY this project exists before touching any code.",
         italic=True)
    doc.add_paragraph()

    h2(doc, "1. What is loan default?")
    para(doc, "A default is when a borrower fails to meet the legal obligations of "
              "their loan — most commonly, they stop paying. It is not a single moment "
              "but a journey of deterioration:")
    bullet(doc, "**DPD (Days Past Due)** — borrower misses a payment (1, 30, 60 days late).")
    bullet(doc, "**Delinquency** — consistently behind (e.g., 30+ DPD).")
    bullet(doc, "**Default** — typically **90+ days past due** (Basel benchmark). Loan assumed unlikely to repay normally.")
    bullet(doc, "**Charge-off / Write-off** — bank gives up collecting and books a loss.")
    para(doc, "In this dataset, bureau_balance.STATUS encodes the journey: "
              "0 = no DPD, 1 = 1-30, 2 = 31-60, 3 = 61-90, 4 = 91-120, 5 = 120+ DPD. "
              "STATUS 3/4/5 is the danger zone.")

    h2(doc, "2. Why do banks lose money due to defaults?")
    para(doc, "A bank's model: borrow cheap (deposits ~3%), lend dear (~9%), pocket the "
              "spread — but the spread must cover defaults. When a loan defaults the bank "
              "loses the principal (depositors' money), not just profit.")
    para(doc, "The killer asymmetry:", bold=True)
    quote(doc, "If a good loan earns ~6% profit and a default loses ~80% of principal, "
               "one default wipes out the profit of roughly 13 good loans.")
    para(doc, "Industry loss vocabulary (memorize):", bold=True)
    table(doc, ["Term", "Meaning", "Example"],
          [["PD", "Probability of Default — how likely (MODEL OUTPUT)", "8%"],
           ["EAD", "Exposure at Default — amount owed at default", "Rs 8 lakh"],
           ["LGD", "Loss Given Default — % unrecoverable after collateral", "60%"]])
    para(doc, "Expected Loss = PD x EAD x LGD.  Your model predicts the PD — the first "
              "and most important term.", bold=True)

    h2(doc, "3. How banks traditionally assess applications — the 5 Cs of Credit")
    table(doc, ["The 5 Cs", "Question", "In the dataset"],
          [["Character", "Will they WANT to repay? (history)", "bureau.csv"],
           ["Capacity", "Can they AFFORD it? (income vs debt)", "AMT_INCOME, AMT_ANNUITY"],
           ["Capital", "Their own skin in the game?", "AMT_DOWN_PAYMENT"],
           ["Collateral", "What can be seized?", "FLAG_OWN_REALTY, AMT_GOODS_PRICE"],
           ["Conditions", "Economy & loan purpose", "NAME_CONTRACT_TYPE, region ratings"]])
    para(doc, "Traditionally assessed by a loan officer using rule-based scorecards, "
              "FICO/bureau scores, and relationship/gut feel.")

    h2(doc, "4. Challenges in manual risk assessment")
    numbered(doc, "**Doesn't scale** — an officer reviews ~20-30 apps/day; Home Credit has 307,511.")
    numbered(doc, "**Inconsistent** — two officers, same file, different decisions.")
    numbered(doc, "**Biased** — unconscious discrimination → fair-lending legal risk (ECOA).")
    numbered(doc, "**Can't handle complexity** — humans juggle ~5-10 variables; the data has hundreds.")
    numbered(doc, "**Slow** — manual underwriting takes days; customers leave.")

    h2(doc, "5. Why machine learning helps")
    para(doc, "ML scales, is consistent, auditable for fairness, handles hundreds of "
              "features, and decides in milliseconds. But the single biggest reason:")
    quote(doc, "ML finds non-linear patterns and INTERACTIONS that humans and simple "
               "scorecards cannot — e.g. 'high income is good UNLESS the person has 5 "
               "recently-opened credit lines AND a history of late payments.'")
    para(doc, "This is why XGBoost (gradient-boosted trees) was chosen and scored "
              "ROC-AUC 0.7864 — trees excel at these interactions.")

    h2(doc, "6. What Home Credit does as a company")
    para(doc, "An international consumer-finance company specialising in lending to "
              "thin-file / underbanked customers in emerging markets — people with little "
              "or no formal credit history, often via point-of-sale loans (financing a "
              "phone, fridge, motorbike in-store).")
    quote(doc, "This is why EXT_SOURCE_1 is 56% missing — many applicants have no "
               "traditional credit score. The modeling challenge is extracting signal from "
               "alternative behavioral data to underwrite people the traditional system rejects.")

    h2(doc, "7. What the TARGET variable means")
    bullet(doc, "**TARGET = 1** — client had payment difficulties (default). The 'bad' class.")
    bullet(doc, "**TARGET = 0** — repaid on time. The 'good' class.")
    bullet(doc, "**Imbalanced: 8.07% default, 91.93% repay.**")
    para(doc, "Consequences of imbalance:", bold=True)
    numbered(doc, "Accuracy is useless — 'predict everyone repays' = 91.93% accurate but catches zero defaulters. Use ROC-AUC.")
    numbered(doc, "Needs special handling — class weights / scale_pos_weight, stratified splits, threshold tuning.")

    h2(doc, "8. Why predicting defaults is valuable")
    numbered(doc, "**Direct loss reduction** — even a 1% lower default rate saves millions on a large book.")
    numbered(doc, "**Risk-based pricing** — price the risk (higher rate) instead of binary approve/reject.")
    numbered(doc, "**Financial inclusion** — safely approve marginal customers (Home Credit's mission).")
    numbered(doc, "**Regulatory capital** — Basel reserves scale with Expected Loss; better PD = freed capital.")
    numbered(doc, "**Speed & experience** — instant decisions win customers and cut cost.")

    divider(doc)
    h2(doc, "Q&A Bank — Session 1 (revise before interviews)")

    h3(doc, "10 Knowledge-Check Questions & Answers")
    qa = [
        ("At how many DPD is a loan typically in default (Basel)?",
         "90 days past due (90 DPD)."),
        ("Expected Loss formula + terms?",
         "EL = PD x EAD x LGD. PD = Probability of Default (model output), "
         "EAD = Exposure at Default, LGD = Loss Given Default (% unrecoverable)."),
        ("How many good loans does one default wipe out?",
         "~13 (80% LGD / 6% profit margin). The asymmetry that justifies risk modeling."),
        ("Name the 5 Cs of Credit.",
         "Character, Capacity, Capital, Collateral, Conditions."),
        ("What does TARGET = 1 mean and what %?",
         "Client had payment difficulties (default). 8.07% of applicants."),
        ("Why is accuracy misleading? Dumb baseline?",
         "Data is imbalanced (8% positive). 'Predict everyone repays' = 91.93% accuracy "
         "but catches zero defaulters. Use ROC-AUC (0.7864)."),
        ("Who does Home Credit lend to?",
         "Thin-file / underbanked customers in emerging markets, often via POS loans."),
        ("Why is EXT_SOURCE_1 56% missing, and why does it matter?",
         "Thin-file applicants have no full bureau score. The bank must rely on "
         "alternative/behavioral data; the missingness itself is a signal."),
        ("What is PD and which part of the project produces it?",
         "Probability of Default; produced by the trained XGBoost model (predict.py / API)."),
        ("One non-linear interaction ML catches that a scorecard can't?",
         "'High income is low-risk UNLESS several recently-opened credit lines AND a "
         "history of late installments' — a conditional interaction."),
    ]
    for i, (q, a) in enumerate(qa, 1):
        para(doc, f"Q{i}. {q}", bold=True)
        para(doc, f"A. {a}")

    h3(doc, "5 Interview Questions & Model Answers")
    iq = [
        ("Walk me through the business problem and why it matters to the bottom line.",
         "Home Credit lends to underbanked customers where ~8% default. Because one "
         "default can wipe out the profit of ~13 good loans, small improvements in "
         "identifying risk save millions. My model predicts each applicant's PD before "
         "approval — driving the approve/reject decision, the interest rate, and "
         "regulatory capital. It turns slow, inconsistent manual review into a scalable, "
         "objective, real-time decision."),
        ("8% default — how does imbalance affect model, metric, and training?",
         "Metric: not accuracy (92% dumb baseline) — I optimise ROC-AUC. Model: XGBoost "
         "with scale_pos_weight / class weighting so the minority class isn't ignored. "
         "Training: stratified splits to preserve the 8% ratio, and threshold tuning for "
         "the business cost trade-off rather than defaulting to 0.5."),
        ("How do you ensure fairness and compliance?",
         "Exclude/audit protected attributes and hunt for proxy discrimination; measure "
         "approval/default rates across groups for disparate impact; use SHAP so every "
         "decision is explainable to a regulator and a declined customer (adverse-action "
         "notices required under ECOA/FCRA). An accurate-but-unfair model is a liability."),
        ("A 20-year loan officer calls your model a black box. Respond.",
         "Agree trust matters and show it isn't a black box. With SHAP I can show, for any "
         "applicant, exactly which factors raised or lowered the score — often the same "
         "reasoning he uses, applied consistently across 300,000 applicants. Position it "
         "as augmenting his judgment, with experts reviewing borderline cases."),
        ("Why does a 1% default reduction matter? Quantify it.",
         "On a billion-dollar book, 1% of defaults avoided — losing ~80% of that "
         "principal — is several million dollars saved annually, dropping almost straight "
         "to the bottom line, and it recurs every cycle."),
    ]
    for i, (q, a) in enumerate(iq, 1):
        para(doc, f"Q{i}. {q}", bold=True)
        para(doc, f"A. {a}")

    h3(doc, "5 Business Case Scenarios & Reasoning")
    cs = [
        ("Threshold dilemma — 15% PD vs 10% cutoff.",
         "Don't hard-reject. Approve-with-conditions: smaller amount, higher rate "
         "(risk-based pricing), collateral/co-signer. A small loan (low EAD) in a growth "
         "region may have acceptable expected loss. Lesson: a probability enables nuanced "
         "pricing, not just approve/reject."),
        ("Recession shift (data/concept drift).",
         "Model trained on a different regime decays and under-predicts risk. Plan: "
         "monitor live vs predicted default rates, watch input drift, retrain on recent "
         "data, add macro features, temporarily tighten threshold. Lesson: models decay; "
         "monitor + retrain."),
        ("Fairness trap — 2x rejection for women.",
         "'The model just uses the data' is not a legal defense — disparate impact is "
         "illegal regardless of intent. Hunt proxy features, remove/neutralise them, apply "
         "fairness constraints, re-validate equity. Accept a small accuracy trade-off for "
         "compliance."),
        ("Cost asymmetry — marketing vs risk.",
         "Frame as False Positives vs False Negatives. Lower threshold = more revenue but "
         "more FN (defaulters let in). Higher = fewer losses but more FP (good customers "
         "rejected). Optimal cutoff is a business decision driven by the cost ratio; the "
         "analyst supplies the curve. Lesson: threshold is economic, not 0.5."),
        ("Thin-file applicant — no bureau data.",
         "Rejecting all no-data applicants defeats Home Credit's mission. Use alternative "
         "data (phone/email/region/prior behavior), graduated lending (small loan to build "
         "history), missing-indicator flags, and manual review for borderline cases. "
         "Lesson: 'no data' is not 'no customer' — it's the underbanked segment."),
    ]
    for i, (q, a) in enumerate(cs, 1):
        para(doc, f"Scenario {i}: {q}", bold=True)
        para(doc, f"-> {a}")

    doc.add_page_break()


def day1_dataset(doc):
    h1(doc, "Day 1 — Session 2: The Dataset (Home Credit Default Risk)")
    para(doc, "Taught from the perspective of a Senior Data Scientist. The dataset is 7 "
              "tables describing one applicant's complete financial life, in two worlds: "
              "INSIDE Home Credit (their own loans) and OUTSIDE (other lenders via the "
              "credit bureau). Everything rolls up to the center table.", italic=True)
    doc.add_paragraph()
    para(doc, "THE core challenge: the center table has ONE row per applicant, but the "
              "other 6 have MANY rows per applicant. All feature engineering = collapsing "
              "many rows -> one row per applicant via aggregation (count, mean, max, sum, ratio).",
         bold=True)

    # Table summary
    h2(doc, "The 7 Tables at a Glance")
    table(doc, ["Table", "Rows", "Key(s)", "What it is"],
          [["application_train", "307,511", "SK_ID_CURR", "The current application + TARGET (fact table)"],
           ["bureau", "1,716,428", "SK_ID_CURR, SK_ID_BUREAU", "Credits at OTHER lenders (credit report)"],
           ["bureau_balance", "27,299,925", "SK_ID_BUREAU", "Monthly status of each external loan"],
           ["previous_application", "1,670,214", "SK_ID_CURR, SK_ID_PREV", "Past applications AT Home Credit"],
           ["installments_payments", "13,605,401", "SK_ID_CURR, SK_ID_PREV", "Repayment records (most behavioral)"],
           ["POS_CASH_balance", "10,001,358", "SK_ID_CURR, SK_ID_PREV", "Monthly POS/cash loan snapshots"],
           ["credit_card_balance", "3,840,312", "SK_ID_CURR, SK_ID_PREV", "Monthly credit card snapshots"]])

    sections = [
        ("1. application_train.csv — The Center of Everything",
         "One row per application (307,511). The fact table — every other table enriches "
         "these rows. PK = SK_ID_CURR. Holds TARGET. The strongest single predictors live "
         "here: EXT_SOURCE_1/2/3 (external credit scores).",
         "Engineered (application.py): CREDIT_INCOME_RATIO (debt vs income), "
         "ANNUITY_INCOME_RATIO (core affordability), CREDIT_TERM_MONTHS, AGE_YEARS, "
         "EMPLOYED_YEARS + IS_UNEMPLOYED (365243 sentinel), EXT_SOURCE_MEAN/MIN/PRODUCT, "
         "and interaction CREDIT_INCOME_x_EXT2 (high debt + low score = compounded risk).",
         "DIRECT — this table holds TARGET; EXT_SOURCE scores are the top predictors."),
        ("2. bureau.csv — Track Record at OTHER Lenders",
         "1,716,428 rows — every credit at other institutions reported to the bureau. Many "
         "rows per applicant. Keys: SK_ID_CURR + SK_ID_BUREAU. This is the credit report "
         "('Character'). Has this person handled credit well elsewhere?",
         "Engineered (bureau.py): BUREAU_LOAN_COUNT, BUREAU_ACTIVE/BAD_DEBT_COUNT, "
         "BUREAU_DEBT_CREDIT_RATIO, BUREAU_OVERDUE_SUM/MAX, BUREAU_DPD_MAX.",
         "STRONG — BAD_DEBT_COUNT>0 (defaulted elsewhere) or high OVERDUE_SUM are powerful flags."),
        ("3. bureau_balance.csv — Month-by-Month History of Each External Loan",
         "27,299,925 rows (2nd biggest) — monthly snapshot of each bureau credit's status. "
         "Key: SK_ID_BUREAU. Reveals the full payment timeline (ever 30/60/90+ late).",
         "Two-step aggregation in bureau.py: (1) bureau_balance -> group by SK_ID_BUREAU "
         "-> per-loan metrics (BB_SEVERE_DPD_MONTHS from STATUS 3/4/5); (2) merge into "
         "bureau -> group by SK_ID_CURR -> BUREAU_BB_SEVERE_DPD_SUM, BUREAU_BB_DPD_RATIO_MEAN.",
         "BEHAVIORAL DEPTH — many SEVERE_DPD months = chronic lateness = future default risk."),
        ("4. previous_application.csv — Past Loans AT Home Credit",
         "1,670,214 rows — every previous application at Home Credit. Keys: SK_ID_CURR + "
         "SK_ID_PREV. Your OWN relationship history: approved or refused before?",
         "Important: NAME_CONTRACT_STATUS (Approved/Refused/Canceled), AMT_APPLICATION vs "
         "AMT_CREDIT (got less than asked?), CODE_REJECT_REASON, CNT_PAYMENT.",
         "Prior refusals/cancellations correlate with risk. CRITICAL: SK_ID_PREV is the "
         "bridge to the 3 behavioral tables below."),
        ("5. installments_payments.csv — The Repayment Receipts (most powerful behavioral table)",
         "13,605,401 rows — actual payment-by-payment record of previous Home Credit loans. "
         "Keys: SK_ID_CURR + SK_ID_PREV. Payment behavior cannot lie — ground truth on repayment.",
         "Derived (installments.py): PAYMENT_DPD = DAYS_ENTRY_PAYMENT - DAYS_INSTALMENT "
         "(positive=late); PAYMENT_DEFICIT = AMT_INSTALMENT - AMT_PAYMENT (positive=underpaid). "
         "-> INSTAL_LATE_RATIO, INSTAL_DPD_MAX/MEAN/STD, INSTAL_UNDERPAYMENT_RATIO, "
         "INSTAL_VERSION_MAX (restructuring = distress). DPD_MAX clipped at 365 (data hygiene).",
         "STRONGEST behavioral link — LATE_RATIO and DPD_MEAN are near-direct proxies for default."),
        ("6. POS_CASH_balance.csv — Monthly Status of POS & Cash Loans",
         "10,001,358 rows — monthly snapshots of previous POS/cash loans. Keys: SK_ID_CURR "
         "+ SK_ID_PREV. Home Credit's core product. Tracks each loan's trajectory.",
         "Important: SK_DPD / SK_DPD_DEF (days past due that month), NAME_CONTRACT_STATUS, "
         "CNT_INSTALMENT_FUTURE (falling steadily = healthy).",
         "Recent rising DPD on existing loans signals current stress -> higher default risk."),
        ("7. credit_card_balance.csv — Monthly Credit Card Behavior",
         "3,840,312 rows — monthly snapshots of previous Home Credit credit cards. Keys: "
         "SK_ID_CURR + SK_ID_PREV. Revolving credit reveals discipline.",
         "Gold signal = UTILIZATION (AMT_BALANCE / AMT_CREDIT_LIMIT_ACTUAL). Maxing out = "
         "stretched even if paying minimum. Cash advances (AMT_DRAWINGS_ATM_CURRENT) = distress.",
         "High utilization + minimum-only payments + cash advances = living beyond means -> risk."),
    ]
    for title, contains, feats, target in sections:
        h2(doc, title)
        para(doc, "What it contains / why: " + contains)
        para(doc, "Features extracted: " + feats)
        p = doc.add_paragraph()
        r = p.add_run("Relationship to TARGET: " + target)
        r.font.italic = True
        r.font.color.rgb = ACCENT

    h2(doc, "How It All Connects — Relational Diagram")
    para(doc, "Three keys hold the universe together: SK_ID_CURR, SK_ID_BUREAU, SK_ID_PREV.")
    mono(doc,
         "                    application_train.csv\n"
         "                    PK: SK_ID_CURR   (holds TARGET)\n"
         "                            |  SK_ID_CURR (1 -> many)\n"
         "      +---------------------+----------------------+\n"
         "      |                     |                      |\n"
         "      v                     v                      v\n"
         "  bureau.csv          previous_application    installments /\n"
         "  PK SK_ID_BUREAU      PK SK_ID_PREV           POS_CASH /\n"
         "  FK SK_ID_CURR        FK SK_ID_CURR           credit_card\n"
         "      |                     |                  FK SK_ID_CURR\n"
         "      | SK_ID_BUREAU        | SK_ID_PREV        & SK_ID_PREV\n"
         "      v                     v\n"
         "  bureau_balance      installments_payments\n"
         "  FK SK_ID_BUREAU     POS_CASH_balance\n"
         "  (monthly rows)      credit_card_balance\n"
         "                      FK SK_ID_PREV")
    bullet(doc, "**SK_ID_CURR** = the customer / this application. The spine; everything rolls up here.")
    bullet(doc, "**SK_ID_BUREAU** = one specific EXTERNAL loan. Links bureau -> bureau_balance.")
    bullet(doc, "**SK_ID_PREV** = one specific past Home Credit loan. Links previous_application -> installments/POS/credit_card.")

    h2(doc, "Worked Example — One Customer (Maria, SK_ID_CURR = 100002)")
    para(doc, "Maria applies for a Rs 600,000 cash loan. Follow her across the tables:")
    bullet(doc, "**application**: AMT_CREDIT=600000, INCOME=200000, EXT_SOURCE_2=0.26 (low), age 25.9, TARGET=1. "
                "-> CREDIT_INCOME_RATIO=3.0, ANNUITY_INCOME_RATIO=0.12. Young + low score + 3x debt-to-income.")
    bullet(doc, "**bureau**: 3 external loans. One Active loan is overdue Rs 4,500. "
                "-> BUREAU_LOAN_COUNT=3, BUREAU_OVERDUE_SUM=4500, DEBT_CREDIT_RATIO~0.68.")
    bullet(doc, "**bureau_balance**: for that loan, 3 months show STATUS 1-2 (late). "
                "-> BUREAU_BB_DPD_MONTHS_SUM=3.")
    bullet(doc, "**previous_application**: 1 prior approved Home Credit loan (SK_ID_PREV=10001), 12 installments.")
    bullet(doc, "**installments**: late on 4 of 12 payments, one underpayment. "
                "-> INSTAL_LATE_RATIO=0.33, INSTAL_DPD_MAX=12. Her own past behavior shows lateness.")
    bullet(doc, "**POS_CASH / credit_card**: monthly SK_DPD ticks up to 12, confirming the late period.")
    para(doc, "Result: ONE wide feature row (~200+ engineered numbers) -> XGBoost -> "
              "PD = 0.71 -> HIGH RISK -> TARGET actually = 1. Every red flag (low EXT_SOURCE, "
              "overdue bureau debt, past lateness) would be INVISIBLE on the application form "
              "alone. That is the entire value of the multi-table dataset: a 360-degree "
              "behavioral portrait from scattered records.", bold=True)

    doc.add_page_break()


def day1_architecture(doc):
    h1(doc, "Day 1 — Session 3: System Architecture (no code)")
    para(doc, "Taught from the perspective of a Senior ML Engineer. The whiteboard talk "
              "for 'walk me through your architecture.' Core principle: SEPARATION OF "
              "CONCERNS — each layer does one job and is independently swappable.", italic=True)
    doc.add_paragraph()
    para(doc, "Two distinct phases (keep them separate): TRAINING (offline, occasional — "
              "turn raw data into a saved model artifact) and INFERENCE (online, per request "
              "— score one new applicant fast). The SAME feature logic runs in both to avoid "
              "train/serve skew.", bold=True)

    h2(doc, "1. The moment raw CSV files arrive")
    para(doc, "Seven raw CSVs land in data/raw/ (307K applications + ~57M supporting rows). "
              "Before any ML:")
    bullet(doc, "**Loading (DataLoader)** — read CSVs into pandas DataFrames. One job; no cleaning.")
    bullet(doc, "**Validation (validator)** — quality gate: expected columns? dtypes? row counts? "
                "Fail loudly here rather than poison the model later. Interview line: 'I validate "
                "data at the boundary and fail fast on bad inputs.'")

    h2(doc, "2. How data flows through the system")
    mono(doc,
         "TRAINING:  Raw CSVs -> Load -> Validate -> Feature Engineering ->\n"
         "           Master Table -> Preprocessing -> Model Training -> Saved Artifact\n\n"
         "INFERENCE: New applicant -> Feature Engineering ->\n"
         "           [Saved Pipeline: Preprocess + Model] -> PD + SHAP -> API / UI")
    para(doc, "One-directional flow. Inference reuses the TAIL of the training pipeline.")

    h2(doc, "3. Feature engineering (the heart of the project)")
    para(doc, "Supporting tables have MANY rows per applicant; the model needs ONE row per "
              "applicant. Feature engineering = collapsing millions of behavioral records into "
              "a few hundred meaningful numbers per person, via aggregation.")
    bullet(doc, "Each table has its own module (bureau.py, installments.py...) that aggregates "
                "to per-applicant metrics (counts, means, maxes, ratios like INSTAL_LATE_RATIO).")
    bullet(doc, "pipeline.py is the ORCHESTRATOR — calls every module and LEFT-JOINs outputs onto "
                "the application table by SK_ID_CURR -> the MASTER TABLE (~300 features).")
    bullet(doc, "Senior move: applicants with no history get NaN after the join; the pipeline fills "
                "those with 0 because 'no history' is a real signal, not missing data.")

    h2(doc, "4. Preprocessing (the last-mile cleaning)")
    para(doc, "Feature engineering creates signal; preprocessing makes it digestible by the "
              "algorithm, consistently for every row:")
    bullet(doc, "Imputation (numeric -> median, categorical -> mode), Scaling, Encoding "
                "(categories -> numbers), Outlier capping.")
    bullet(doc, "Wrapped in a scikit-learn Pipeline TOGETHER with the model, saved as ONE artifact "
                "-> identical transforms in training and inference -> no skew, ever.")
    para(doc, "FE vs preprocessing: FE is domain-driven creativity (inventing INSTAL_LATE_RATIO); "
              "preprocessing is mechanical algorithm-driven cleaning (impute/scale/encode). "
              "First needs banking knowledge, second needs ML knowledge.", bold=True)

    h2(doc, "5. Model training")
    para(doc, "Learning the patterns separating defaulters from repayers — a disciplined process, "
              "not just .fit():")
    bullet(doc, "5-fold cross-validation — honest, stable performance estimate (not a lucky split).")
    bullet(doc, "Model comparison — XGBoost, LightGBM, etc., picked by ROC-AUC.")
    bullet(doc, "Hyperparameter tuning with Optuna.")
    bullet(doc, "Final fit on all data -> serialize whole pipeline to models/v2/pipeline.pkl + "
                "metadata.json (AUC, threshold, training date).")

    h2(doc, "6. SHAP explainability (the trust layer)")
    para(doc, "A high-AUC black box is useless to a bank — regulators/officers won't trust it and "
              "the law (ECOA/FCRA) requires telling a rejected customer WHY. SHAP (Shapley values) "
              "answers: for THIS applicant, how much did each feature push the prediction up/down.")
    bullet(doc, "Global — which features matter most across all customers.")
    bullet(doc, "Local (waterfall) — for one applicant: 'high-risk mainly due to low EXT_SOURCE, "
                "overdue bureau debt, past late payments.'")
    para(doc, "Note: global SHAP works; wiring the per-applicant waterfall into predict.py is the "
              "one remaining to-do item.", italic=True)

    h2(doc, "7. FastAPI (serving layer)")
    para(doc, "Turns the saved model into a live web service other software can call. Endpoints: "
              "POST /predict (applicant data -> PD + explanation), GET /health, GET /explain. "
              "Validates inputs via Pydantic; async, production-grade, auto-documented. "
              "Serves MACHINES (e.g., a loan-origination system) over HTTP in milliseconds.")

    h2(doc, "8. Streamlit (UI layer)")
    para(doc, "Human-friendly dashboard — the demo face. Sidebar customer input, risk score, SHAP "
              "waterfall, model-performance and business-context tabs. FastAPI serves machines; "
              "Streamlit serves humans (analysts, stakeholders, interviewers). Makes the portfolio "
              "project clickable and impressive.")

    h2(doc, "9. Docker (reproducibility & deployment)")
    para(doc, "Solves 'it works on my machine.' Packages code + Python version + all libraries + "
              "model into a container that runs identically anywhere. docker-compose orchestrates "
              "the API + UI together with one command. The difference between a notebook and a "
              "deployable product.")

    h2(doc, "10. How the entire system works together — end-to-end diagram")
    mono(doc,
         "TRAINING PHASE (offline)\n"
         " (1) DATA: 7 raw CSVs -> DataLoader -> Validator\n"
         " (2) FEATURE ENGINEERING: application/bureau/previous/installments/\n"
         "     pos_cash/credit_card -> pipeline.py = MASTER TABLE (~300 feats)\n"
         " (3) PREPROCESSING: impute + scale + encode + cap (in sklearn Pipeline)\n"
         " (4) MODELING: 5-fold CV -> model compare -> Optuna -> final fit\n"
         "     evaluate.py (ROC-AUC, threshold) | explain.py (SHAP)\n"
         "         => SAVED ARTIFACT: models/v2/pipeline.pkl + metadata.json\n"
         "                         |  (load once at startup)\n"
         "INFERENCE PHASE (online, per request)\n"
         "   New applicant -> same Feature Engineering ->\n"
         "        [Pipeline: preprocess + model] -> PD + SHAP\n"
         "          |                                  |\n"
         "          v                                  v\n"
         "   (5) FastAPI (machines)            (6) Streamlit (humans)\n"
         "       /predict /health /explain         input + score + SHAP plot\n\n"
         " (7) DOCKER wraps EVERYTHING -> portable, reproducible, deployable\n"
         "     (docker-compose runs API + UI together)")
    para(doc, "One-liner to tie it together: 'A layered system with clean separation of concerns. "
              "The same feature pipeline runs in training and inference to eliminate skew; the model "
              "and its preprocessing travel together as one serializable artifact; SHAP provides the "
              "trust layer; Docker makes it portable. That is the difference between a Kaggle "
              "notebook and a deployable product.'", bold=True)

    doc.add_page_break()


def day1_summary(doc):
    h1(doc, "Day 1 — Complete Summary (revision page)")
    para(doc, "One-page consolidation of Day 1's three sessions plus the key doubts cleared. "
              "Read this first when revising.", italic=True)
    doc.add_paragraph()

    h2(doc, "The Project in One Sentence")
    para(doc, "For a new loan applicant, predict the PROBABILITY they will default — so the bank "
              "can approve, reject, or adjust terms BEFORE lending. (A weather forecast for a loan: "
              "not 'will rain' but '70% chance of rain.')", bold=True)

    h2(doc, "Session 1 — The Business Problem")
    bullet(doc, "Default = borrower stops paying; a journey: missed payment -> DPD -> delinquency -> "
                "**default (90+ DPD, Basel)** -> charge-off.")
    bullet(doc, "Banks borrow cheap (~3%), lend dear (~9%); the spread must cover losses. "
                "**One default wipes out ~13 good loans** — the asymmetry that justifies risk modeling.")
    bullet(doc, "**Expected Loss = PD x EAD x LGD.** PD (model output) = Probability of Default; "
                "EAD = Exposure at Default; LGD = Loss Given Default.")
    bullet(doc, "Home Credit lends to thin-file/underbanked customers -> EXT_SOURCE_1 is 56% missing "
                "-> challenge = extract signal from alternative behavioral data.")
    bullet(doc, "TARGET: 1 = default (8.07%), 0 = repaid (91.93%). **Imbalanced -> use ROC-AUC, not accuracy.**")

    h2(doc, "Session 2 — The Dataset")
    bullet(doc, "7 tables, two worlds: INSIDE Home Credit (application, previous, installments, "
                "POS_CASH, credit_card) and OUTSIDE via bureau (bureau, bureau_balance).")
    bullet(doc, "Core challenge: history tables have MANY rows per person; model needs ONE row per person.")
    bullet(doc, "3 join keys: **SK_ID_CURR** (the person/spine), **SK_ID_BUREAU** (one external loan), "
                "**SK_ID_PREV** (one past Home Credit loan).")
    bullet(doc, "Result = MASTER TABLE. Verified on disk: **307,511 rows x 246 columns** "
                "('one row per person, ~245 features').")

    h2(doc, "Session 3 — System Architecture")
    bullet(doc, "Core principle: **separation of concerns** — each layer/file does one job, swappable.")
    bullet(doc, "Layers: Data (load+validate) -> Feature Engineering (aggregate+join) -> Preprocessing "
                "-> Modeling -> Serving (FastAPI) + UI (Streamlit), all wrapped in Docker.")
    bullet(doc, "**Two journeys:** TRAINING (once) builds & saves pipeline.pkl 'brain'; "
                "INFERENCE (per applicant) loads the brain -> PD + SHAP -> API/UI.")
    bullet(doc, "SHAP = trust layer (explains WHY). Verified report card (metadata.json): "
                "**XGBoost won, Test ROC-AUC 0.7864**, beating LightGBM, Logistic Regression, Random Forest.")

    h2(doc, "Key Clarifications (the gotchas)")
    numbered(doc, "**Feature Engineering != filling nulls.** FE = creating new signal (INSTAL_LATE_RATIO); "
                  "filling nulls = preprocessing (imputation), done later.")
    numbered(doc, "**Aggregate FIRST, then join.** Can't join history tables directly (too many rows); "
                  "squash each to 1 row/person, then join.")
    numbered(doc, "**Validation** = gatekeeper check after loading (columns/types/row-counts); 'fail fast'. "
                  "Different from preprocessing (cleaning) and cross-validation (checking the model).")
    numbered(doc, "**Code = chain of calls.** Orchestrator (pipeline.py, train.py) calls worker files "
                  "(bureau.py, loader.py); each does one job and returns. Boss coordinates, never does grunt work.")
    numbered(doc, "**Same FE in training AND inference** -> bundling preprocess+model in one .pkl prevents "
                  "train/serve skew.")

    h2(doc, "The Corrected Mental Flow")
    mono(doc,
         "1. Load 7 tables\n"
         "2. AGGREGATE each history table -> 1 row/person   |\n"
         "3. Engineer new columns (ratios, counts, flags)   | FEATURE ENGINEERING\n"
         "4. JOIN all -> master table (307k x 246)          |\n"
         "5. Fill nulls, scale, encode  -------------------- PREPROCESSING\n"
         "6. Train (CV + compare + tune) -> save pipeline.pkl  MODELING\n"
         "7. Serve: FastAPI (machines) / Streamlit (humans), wrapped in Docker")

    h2(doc, "Files to Inspect on Disk (Day 1)")
    bullet(doc, "data/processed/master_train.csv — the master table (scroll right for BUREAU_*, INSTAL_*, etc.).")
    bullet(doc, "models/v2/metadata.json — the model report card (0.7864, 8% default rate, model comparison).")
    bullet(doc, "outputs/plots/shap_global_importance.png — explainability as a picture.")
    bullet(doc, "data/raw/application_train.csv (1 row/person) vs installments_payments.csv (many rows/person).")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# BUILD
# ─────────────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup_base_styles(doc)
    add_title_page(doc)
    day1_summary(doc)
    day1_business(doc)
    day1_dataset(doc)
    day1_architecture(doc)
    try:
        doc.save(OUT_PATH)
        print(f"Study guide written to: {OUT_PATH}")
    except PermissionError:
        # File is open in Word — save to a fallback name so work isn't lost.
        from datetime import datetime
        fallback = OUT_PATH.with_name(
            f"Interview_Prep_Study_Guide_{datetime.now():%H%M%S}.docx"
        )
        doc.save(fallback)
        print(f"MAIN FILE LOCKED (open in Word). Saved to fallback: {fallback}")
        print("Close the main .docx, delete the fallback, and re-run to consolidate.")


if __name__ == "__main__":
    build()
